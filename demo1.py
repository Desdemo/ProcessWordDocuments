import sys
import importlib
from docx import Document
from calendary import Calendary

def main():
    # 创建文档对象
    document = Document('G:\demo1.docx')
    # 读取文档中所有的段落列表
    # ps = document.paragraphs
    # # 每个段落有两个属性：style和text
    # ps_detail = [(x.text, x.style.name) for x in ps]
    # with open('out.tmp', 'w+') as fout:
    #     fout.write('')
    # # 读取段落并写入一个文件
    # with open('out.tmp', 'a+') as fout:
    #     for p in ps_detail:
    #         fout.write(p[0] + '\t' + p[1] + '\n\n')

    # 读取文档中的所有段落的列表
    tables = document.tables
    #获取工作日
    cal = Calendary(2018)
    date_list = []
    cal = cal.month(8, work=True, workweek_start=1, workweek_end=5)
    for weekday, date in cal:
        date_list.append(str(date))
    for day in date_list:
        tables[0].cell(0, 1).text = "日期:" + day
        name = day +'-XXX'+'.docx'
        document.save(name)
    # 遍历table，并将所有单元格内容写入文件中
    # with open('out.tmp', 'a+') as fout:
    #     for table in tables:
    #         for row in table.rows:
    #             for cell in row.cells:
    #                 fout.write(cell.text + '\t')
    #             fout.write('\n')
#
if __name__ == '__main__':
    main()